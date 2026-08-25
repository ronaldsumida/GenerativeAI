#!/usr/bin/env python3
"""Render a structured travel itinerary to an editable DOCX plus HTML preview."""

import hashlib
import html
import json
import os
import re
import sys
import traceback
import uuid

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "static", "itineraries")

DATE_FILL = "E0ECF8"
ACTIVITY_FILL = "FFFFE0"
HOTEL_FILL = "E0FFE0"
TRANSPORT_FILLS = ["FFF0FF", "E6E6FF", "FFEAD9", "E8F4FF"]
BORDER_COLOR = "666666"

# Defense in depth: the skill is instructed not to send pricing information,
# but the renderer also strips obvious pricing lines if one slips through.
_CURRENCY_RE = re.compile(
    r"(?:[$€£¥]\s?\d|\b(?:USD|EUR|GBP|JPY|CAD|AUD|CHF|CNY|RMB|HKD|AED|NPR|INR)\s*\d)",
    re.IGNORECASE,
)
_PRICE_LABEL_RE = re.compile(
    r"\b(?:price|fare|room rate|nightly rate|tax(?:es)?|fee(?:s)?|total|amount paid|balance due|cost)\b",
    re.IGNORECASE,
)


def _is_pricing_line(value):
    text = str(value or "").strip()
    if not text:
        return False
    return bool(_CURRENCY_RE.search(text) or (_PRICE_LABEL_RE.search(text) and re.search(r"\d", text)))


def _clean_lines(lines):
    return [str(line) for line in (lines or []) if not _is_pricing_line(line)]


def _clean_item(item):
    item = dict(item or {})
    if _is_pricing_line(item.get("title")):
        item["title"] = ""
    item["lines"] = _clean_lines(item.get("lines"))
    return item


def _set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=70, start=85, bottom=70, end=85):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), BORDER_COLOR)


def _set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_repeat_table_widths(table):
    widths = [Inches(1.685), Inches(3.877), Inches(1.938)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    return widths


def _style_paragraph(paragraph, centered=False, keep_with_next=False):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_with_next = keep_with_next
    if centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_line(paragraph, text, bold=False):
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    return run


def _add_item(cell, item, blank_before=False):
    item = _clean_item(item)
    title = item.get("title") or ""
    lines = item.get("lines") or []
    if not title and not lines:
        return

    if blank_before and cell.paragraphs and cell.paragraphs[-1].text:
        p = cell.add_paragraph()
        _style_paragraph(p)

    p = cell.paragraphs[-1] if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text else cell.add_paragraph()
    _style_paragraph(p)
    if title:
        _add_line(p, title, bold=True)
    for line in lines:
        if p.text:
            p.add_run("\n")
        _add_line(p, line)


def _transport_fill(items):
    key = "|".join((item.get("title") or "") for item in items)
    digest = hashlib.sha1(key.encode("utf-8")).digest()[0] if key else 0
    return TRANSPORT_FILLS[digest % len(TRANSPORT_FILLS)]


def _render_day(table, day):
    date_row = table.add_row()
    merged = date_row.cells[0].merge(date_row.cells[1]).merge(date_row.cells[2])
    merged.text = ""
    p = merged.paragraphs[0]
    _style_paragraph(p, centered=True, keep_with_next=True)
    run = p.add_run(day.get("date") or "")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    _set_cell_shading(merged, DATE_FILL)
    _set_cell_margins(merged, top=20, bottom=20, start=60, end=60)
    _set_cell_borders(merged)

    row = table.add_row()
    _set_row_cant_split(row)
    cells = row.cells
    for cell in cells:
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        _set_cell_margins(cell)
        _set_cell_borders(cell)
        _style_paragraph(cell.paragraphs[0])

    transportation = [_clean_item(i) for i in (day.get("transportation") or [])]
    transportation = [i for i in transportation if i.get("title") or i.get("lines")]
    for i, item in enumerate(transportation):
        _add_item(cells[0], item, blank_before=(i > 0))
    if transportation:
        _set_cell_shading(cells[0], _transport_fill(transportation))

    plans = [_clean_item(i) for i in (day.get("plans") or [])]
    plans = [i for i in plans if i.get("title") or i.get("lines")]
    for i, item in enumerate(plans):
        _add_item(cells[1], item, blank_before=(i > 0))
    if plans and all((item.get("kind") == "activity") for item in plans):
        _set_cell_shading(cells[1], ACTIVITY_FILL)

    lodging = _clean_item(day.get("lodging") or {})
    if lodging.get("title") or lodging.get("lines"):
        title = lodging.get("title") or ""
        lines = lodging.get("lines") or []
        p = cells[2].paragraphs[0]
        if title:
            _add_line(p, title, bold=bool(lodging.get("full_details")))
        for line in lines:
            if p.text:
                p.add_run("\n")
            _add_line(p, line)
        if lodging.get("full_details"):
            _set_cell_shading(cells[2], HOTEL_FILL)


def _html_item(item):
    item = _clean_item(item)
    title = item.get("title") or ""
    lines = item.get("lines") or []
    if not title and not lines:
        return ""
    parts = []
    if title:
        parts.append(f"<strong>{html.escape(title)}</strong>")
    parts.extend(html.escape(str(line)) for line in lines)
    return "<div class=\"preview-item\">" + "<br>".join(parts) + "</div>"


def _build_preview_html(data):
    rows = []
    for day in data.get("days") or []:
        date_text = html.escape(str(day.get("date") or ""))
        rows.append(f'<tr class="date-row"><td colspan="3">{date_text}</td></tr>')

        transportation = [_clean_item(i) for i in (day.get("transportation") or [])]
        transportation = [i for i in transportation if i.get("title") or i.get("lines")]
        t_html = "".join(_html_item(i) for i in transportation)
        t_style = f' style="background:#{_transport_fill(transportation)}"' if transportation else ""

        plans = [_clean_item(i) for i in (day.get("plans") or [])]
        plans = [i for i in plans if i.get("title") or i.get("lines")]
        p_html = "".join(_html_item(i) for i in plans)
        p_style = ' style="background:#FFFFE0"' if plans and all(i.get("kind") == "activity" for i in plans) else ""

        lodging = _clean_item(day.get("lodging") or {})
        l_html = _html_item(lodging)
        l_style = ' style="background:#E0FFE0"' if lodging.get("full_details") and l_html else ""

        rows.append(
            "<tr class=\"content-row\">"
            f"<td class=\"transport\"{t_style}>{t_html}</td>"
            f"<td class=\"plans\"{p_style}>{p_html}</td>"
            f"<td class=\"lodging\"{l_style}>{l_html}</td>"
            "</tr>"
        )

    title = html.escape(str(data.get("title") or "Travel Itinerary"))
    return f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{title}</title>
<style>
html,body{{margin:0;background:#f3f3f3;font-family:Calibri,Arial,sans-serif;color:#111}}
.page{{width:min(8.5in,calc(100% - 28px));margin:18px auto;background:#fff;box-shadow:0 2px 14px rgba(0,0,0,.18);padding:.5in;min-height:11in;box-sizing:border-box}}
table{{border-collapse:collapse;width:100%;table-layout:fixed;font-size:13px;line-height:1.22}}
col.transport{{width:22.5%}} col.plans{{width:51.7%}} col.lodging{{width:25.8%}}
td{{border:1px solid #666;padding:7px 8px;vertical-align:top;overflow-wrap:anywhere}}
.date-row td{{background:#E0ECF8;text-align:center;padding:3px 6px;font-size:12px}}
.preview-item+.preview-item{{margin-top:14px}}
strong{{font-weight:700}}
@media(max-width:700px){{.page{{width:100%;margin:0;padding:14px;box-shadow:none;min-height:0}}table{{font-size:11px}}td{{padding:5px}}}}
</style>
</head>
<body>
<div class="page">
<table>
<colgroup><col class="transport"><col class="plans"><col class="lodging"></colgroup>
<tbody>{''.join(rows)}</tbody>
</table>
</div>
</body>
</html>"""


def build_itinerary(data):
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    stem = f"travel-itinerary-{uuid.uuid4().hex}"
    filename = f"{stem}.docx"
    path = os.path.join(OUTPUT_DIR, filename)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)

    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for day in data.get("days") or []:
        _render_day(table, day)

    _set_repeat_table_widths(table)

    if doc.paragraphs:
        doc.paragraphs[-1].paragraph_format.space_before = Pt(0)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(0)
        doc.paragraphs[-1].paragraph_format.line_spacing = 0.1

    doc.save(path)

    preview_path = os.path.join(OUTPUT_DIR, f"{stem}.html")
    with open(preview_path, "w", encoding="utf-8") as f:
        f.write(_build_preview_html(data))

    return path


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "expected one JSON argument"}))
        sys.exit(1)

    try:
        data = json.loads(sys.argv[1])
        if not isinstance(data, dict) or not isinstance(data.get("days"), list):
            raise ValueError("JSON must be an object containing a 'days' array")
        path = build_itinerary(data)
        print(json.dumps({"filename": os.path.basename(path)}))
    except Exception as e:
        print(json.dumps({
            "error": f"failed to render itinerary: {e}",
            "traceback": traceback.format_exc(),
        }))
        sys.exit(1)


if __name__ == "__main__":
    main()
